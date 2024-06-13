import os
import pandas as pd
from datetime import datetime


class DataObject:
    def __init__(self, data):
        self.data = data

    '''
    def apply_gpt4(self, column_name):
        if column_name in self.data.columns:
            self.data[column_name] = self.data[column_name].apply(process_with_gpt4)
    '''

import os
import pandas as pd
import json
import chardet
from docx import Document
from bs4 import BeautifulSoup

def read_data(file_path):
    extension = os.path.splitext(file_path)[1]
    data_object = None

    try:
        if extension == '.py':
            with open(file_path, 'r') as f:
                data_object = {'script': f.read()}
        elif extension == '.csv':
            try:
                data_object = pd.read_csv(file_path, error_bad_lines=False).to_dict()
            except pd.errors.ParserError as e:
                print(f"An error occurred while processing {file_path}: {str(e)}")
                return None
        elif extension == '.txt':
            with open(file_path, 'r') as f:
                data_object = {'text': f.readlines()}
        elif extension == '.json':
            with open(file_path, 'r') as f:
                data_object = json.load(f)
        elif extension == '.xlsx':
            data_object = pd.read_excel(file_path).to_dict()
        elif extension == '.sql':
            with open(file_path, 'r') as f:
                data_object = {'sql': f.read()}
        elif extension == '.docx':
            doc = Document(file_path)
            data_object = {'text': [p.text for p in doc.paragraphs]}
        elif extension == '.html':
            rawdata = open(file_path, "rb").read()
            result = chardet.detect(rawdata)
            encoding = result['encoding']

            with open(file_path, 'r', encoding=encoding) as f:
                soup = BeautifulSoup(f, 'html.parser')
                data_object = {'html': str(soup)}
    except Exception as e:
        print(f"An error occurred while processing {file_path}: {str(e)}")
        return None

    return data_object




import os
import pandas as pd
import json
from openpyxl import Workbook

def write_to_excel(data_objects, output_folder):
    try:
        if not os.path.exists(output_folder):
            os.makedirs(output_folder)

        wb = Workbook()
        for file_name, data_object in data_objects.items():
            ws = wb.create_sheet(title=file_name[:30])
            for key, value in data_object.items():
                ws.append([key, value])

        output_path = os.path.join(output_folder, 'combined_data.xlsx')
        wb.save(output_path)
        print(f"Excel file successfully saved at {output_path}")
    except Exception as e:
        print(f"An error occurred while writing to Excel: {str(e)}")

def write_to_csv(data_object, output_folder, file_name):
    try:
        if not os.path.exists(output_folder):
            os.makedirs(output_folder)

        output_path = os.path.join(output_folder, f"{file_name}.csv")
        pd.DataFrame(data_object).to_csv(output_path, index=False)
        print(f"CSV file successfully saved at {output_path}")
    except Exception as e:
        print(f"An error occurred while writing to CSV: {str(e)}")

def write_to_json(data_object, output_folder, file_name):
    try:
        if not os.path.exists(output_folder):
            os.makedirs(output_folder)

        output_path = os.path.join(output_folder, f"{file_name}.json")
        with open(output_path, 'w') as f:
            json.dump(data_object, f)
        print(f"JSON file successfully saved at {output_path}")
    except Exception as e:
        print(f"An error occurred while writing to JSON: {str(e)}")

def write_to_txt(data_object, output_folder, file_name):
    try:
        if not os.path.exists(output_folder):
            os.makedirs(output_folder)

        output_path = os.path.join(output_folder, f"{file_name}.txt")
        with open(output_path, 'w') as f:
            f.write(str(data_object))
        print(f"Text file successfully saved at {output_path}")
    except Exception as e:
        print(f"An error occurred while writing to Text: {str(e)}")


from docx import Document
import pandas as pd
from datetime import datetime
import os

def write_processed_data(data_object, output_folder, filename):
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

import logging

# Initialize logging
logging.basicConfig(filename='data_processing.log', level=logging.INFO)

def write_processed_data(df, output_folder, base_filename):
    try:
        # Write the DataFrame to a CSV file
        csv_output_path = os.path.join(output_folder, f"{base_filename}.csv")
        df.to_csv(csv_output_path, index=False)
        logging.info(f"Successfully wrote CSV file to {csv_output_path}")
    except Exception as e:
        logging.error(f"An error occurred while writing CSV file: {str(e)}")

    try:
        # Write the DataFrame to an Excel file
        excel_output_path = os.path.join(output_folder, f"{base_filename}.xlsx")
        df.to_excel(excel_output_path, index=False)
        logging.info(f"Successfully wrote Excel file to {excel_output_path}")
    except Exception as e:
        logging.error(f"An error occurred while writing Excel file: {str(e)}")

    try:
        # Write the DataFrame to a Word document
        doc_output_path = os.path.join(output_folder, f"{base_filename}.docx")
        doc = Document()
        if 'Key' in df.columns and 'Value' in df.columns:
            for index, row in df.iterrows():
                doc.add_paragraph(f"{row['Key']}: {row['Value']}")
            doc.save(doc_output_path)
            logging.info(f"Successfully wrote Word document to {doc_output_path}")
        else:
            logging.warning(f"DataFrame does not contain 'Key' and 'Value' columns. Skipping Word document creation.")
    except Exception as e:
        logging.error(f"An error occurred while writing Word document: {str(e)}")

    try:
        # Write the DataFrame to an HTML document
        html_output_path = os.path.join(output_folder, f"{base_filename}.html")
        df.to_html(html_output_path, index=False)
        logging.info(f"Successfully wrote HTML document to {html_output_path}")
    except Exception as e:
        logging.error(f"An error occurred while writing HTML document: {str(e)}")



def read_local(root_folder, output_folder):
    data_objects = {}
    print(f"Initiating data processing from root folder: {root_folder}")

    for subdir, _, files in os.walk(root_folder):
        print(f"Scanning directory: {subdir}")

        for filename in files:
            file_path = os.path.join(subdir, filename)
            print(f"Reading file: {file_path}")

            data_object = read_data(file_path)
            if data_object:
                print(f"Successfully read data from file: {filename}")
                data_objects[filename] = data_object

                print(f"Initiating write operation for file: {filename}")
                write_processed_data(data_object, output_folder, filename)
                print(f"Successfully wrote processed data for file: {filename}")

    print("Commencing final Excel workbook compilation.")
    # write_to_excel(data_objects, output_folder)  # Uncomment if you still want to write to Excel
    print("Data processing pipeline has reached completion.")

if __name__ == "__main__":
    root_folder = os.getcwd()
    output_folder = os.path.join(root_folder, "playground\\data")
    print("Data processing pipeline initiated.")
    read_local(root_folder, output_folder)
    print("Data processing pipeline terminated.")
