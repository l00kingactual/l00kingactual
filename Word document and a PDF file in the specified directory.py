import os
import glob
import datetime
import logging
from pathlib import Path
from docx import Document
from docx2pdf import convert
from PyPDF2 import PdfMerger

# Setup logging
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')

# Directory path
directory = r'C:\\Users\\actua\\OneDrive\\Desktop\\ideas\\manchester'

# Ensure the output directory exists
combined_folder = Path(directory) / 'combined_output'
combined_folder.mkdir(parents=True, exist_ok=True)

# Create a new Word document
doc = Document()

# Get all files except ZIP files
file_paths = glob.glob(os.path.join(directory, '*'))
files = [file for file in file_paths if not file.endswith('.zip')]

# Sort files by last modified date
files.sort(key=lambda x: os.path.getmtime(x))

# Define a function to read the file content with error handling
def read_file_content(file_path):
    encodings = ['utf-8', 'latin-1', 'cp1252']
    for enc in encodings:
        try:
            with open(file_path, 'r', encoding=enc) as f:
                logging.debug(f"Reading file {file_path} with encoding {enc}")
                return f.read()
        except UnicodeDecodeError as e:
            logging.error(f"UnicodeDecodeError for file {file_path} with encoding {enc}: {e}")
        except Exception as e:
            logging.error(f"Error reading file {file_path} with encoding {enc}: {e}")
    return None

# Define a function to sanitize the content
def sanitize_content(content):
    return ''.join(char for char in content if ord(char) > 31 or char in ('\n', '\r'))

pdf_files = []

try:
    for file in files:
        try:
            file_name = os.path.basename(file)
            mod_time = os.path.getmtime(file)
            date_time = datetime.datetime.fromtimestamp(mod_time).strftime('%Y-%m-%d %H:%M:%S')

            if file_name.endswith('.docx'):
                # Add date time as heading 1
                doc.add_heading(date_time, level=1)
                # Add file name as heading 2
                doc.add_heading(file_name, level=2)

                # Read file contents and add to the document
                content = read_file_content(file)
                if content:
                    sanitized_content = sanitize_content(content)
                    doc.add_paragraph(sanitized_content)
                else:
                    doc.add_paragraph("Could not read the content of this file due to encoding issues.")
                
                # Add a page break after each file
                doc.add_page_break()

                logging.info(f"Processed file: {file_name}")

                # Convert Word document to PDF
                temp_pdf_path = combined_folder / (file_name + '.pdf')
                convert(file, temp_pdf_path)
                pdf_files.append(temp_pdf_path)

            elif file_name.endswith('.pdf'):
                pdf_files.append(file)
                logging.info(f"Processed PDF file: {file_name}")

        except Exception as e:
            logging.error(f"Error processing file {file}: {e}")

    # Save the Word document
    output_path_docx = os.path.join(directory, 'combined_document.docx')
    doc.save(output_path_docx)
    logging.info(f"Combined Word document saved as {output_path_docx}")

    # Convert the combined Word document to PDF
    output_path_pdf = os.path.join(directory, 'combined_document.pdf')
    convert(output_path_docx, output_path_pdf)
    pdf_files.append(output_path_pdf)
    logging.info(f"Combined PDF document saved as {output_path_pdf}")

    # Combine all PDFs into one
    merger = PdfMerger()
    for pdf_file in pdf_files:
        merger.append(str(pdf_file))

    final_pdf_output = combined_folder / 'final_combined_document.pdf'
    merger.write(str(final_pdf_output))
    merger.close()
    logging.info(f"All PDFs combined into {final_pdf_output}")

except Exception as e:
    logging.critical(f"Critical error in processing: {e}")
